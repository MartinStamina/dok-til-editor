"""dok-til-editor: PDF/DOCX -> CKEditor 4-vennlig HTML.

Regler dokumentert i regler.md. Endringer her skal speile en regel-endring.
"""
from __future__ import annotations

import argparse
import base64
import mimetypes
import re
import sys
from collections import Counter
from pathlib import Path

ALLOWED_TAGS = [
    "h2", "h3", "h4",
    "p", "br",
    "ul", "ol", "li",
    "strong", "em",
    "a",
    "img",
    "table", "thead", "tbody", "tr", "th", "td",
]
ALLOWED_ATTRS = {
    "a": ["href"],
    "img": ["src", "alt", "width", "height"],
    "h2": ["id"], "h3": ["id"], "h4": ["id"],
}
ALLOWED_PROTOCOLS = ["http", "https", "mailto", "data"]

NOISE_PATTERNS = [
    re.compile(r"^\s*\d+\s*$"),
    re.compile(r"^\s*sivu\s+\d+\s*[/|]\s*\d+\s*$", re.I),
    re.compile(r"^\s*page\s+\d+\s+of\s+\d+\s*$", re.I),
]

MIN_HEADING_LEN = 3
MAX_HEADING_LEN = 80
HEADING_DEDUPE_WINDOW = 50
HEADING_MERGE_MAX_LEN = 40
SCANNED_PDF_CHARS_PER_PAGE = 200
PARAGRAPH_GAP_MULTIPLIER = 1.6
SENTENCE_END = re.compile(r"[.!?:;»\")\]]\s*$")
HYPHEN_END = re.compile(r"[\-­‐‑]\s*$")
BULLET_CHARS = set("•·‧●▪◦◆▶▸")
BULLET_SPLIT = re.compile(r"\s*[" + re.escape("".join(BULLET_CHARS)) + r"]\s+")


def image_placeholder(page: int | None = None, filename: str | None = None) -> str:
    page_part = f" — page {page}" if page else ""
    if filename:
        return f"[ IMAGE — attachment: {filename}{page_part} ]"
    return f"[ MISSING IMAGE — INSERT IMAGE AND DESCRIPTION HERE{page_part} ]"


def image_element(page: int | None, filename: str | None, image_dir: Path | None, embed: bool = False) -> tuple[str, str]:
    """Returner (tag, content) for et bilde-element.

    Hvis embed=True og filen eksisterer: returner ("__raw__", "<p><img src='data:...;base64,...'></p>").
    Ellers: returner ("p", placeholder-tekst som peker til ekstrahert fil).
    """
    if embed and filename and image_dir:
        full = image_dir / filename
        if full.exists() and full.stat().st_size > 0:
            mime, _ = mimetypes.guess_type(str(full))
            if not mime:
                mime = "image/jpeg"
            data = base64.b64encode(full.read_bytes()).decode("ascii")
            alt = f"page {page} image" if page else "image"
            html_img = f'<p><img src="data:{mime};base64,{data}" alt="{alt}"></p>'
            return ("__raw__", html_img)
    return ("p", image_placeholder(page, filename))


def extract_pdf_images(pdf_path: Path, out_dir: Path, min_size: int = 50) -> list[dict]:
    """Ekstrah bilder fra PDF. Returnerer liste av dicts med page, filename, top.

    Filtrerer ut bilder mindre enn min_size x min_size (logoer/ornamenter).
    """
    import fitz

    out_dir.mkdir(parents=True, exist_ok=True)
    results: list[dict] = []
    doc = fitz.open(str(pdf_path))
    try:
        for page_num, page in enumerate(doc, start=1):
            images_on_page = page.get_images(full=True)
            for idx, img_info in enumerate(images_on_page, start=1):
                xref = img_info[0]
                try:
                    base = doc.extract_image(xref)
                except Exception:
                    continue
                width = base.get("width", 0)
                height = base.get("height", 0)
                if width < min_size or height < min_size:
                    continue
                ext = base.get("ext", "png")
                image_bytes = base["image"]
                filename = f"page{page_num}_img{idx}.{ext}"
                (out_dir / filename).write_bytes(image_bytes)
                bbox_rects = page.get_image_rects(xref)
                top = float(bbox_rects[0].y0) if bbox_rects else 0.0
                results.append({"page": page_num, "filename": filename, "top": top, "width": width, "height": height})
    finally:
        doc.close()
    return results


def extract_docx_images(docx_path: Path, out_dir: Path) -> list[dict]:
    """Ekstrah inline-bilder fra DOCX. Returnerer liste av dicts med filename."""
    import docx

    out_dir.mkdir(parents=True, exist_ok=True)
    results: list[dict] = []
    doc = docx.Document(str(docx_path))
    image_parts = [p for p in doc.part.related_parts.values() if p.content_type.startswith("image/")]
    for idx, part in enumerate(image_parts, start=1):
        ext = part.content_type.split("/")[-1].split("+")[0]
        if ext == "jpeg":
            ext = "jpg"
        filename = f"img{idx}.{ext}"
        (out_dir / filename).write_bytes(part.blob)
        results.append({"filename": filename})
    return results


def is_noise(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    return any(p.match(line) for p in NOISE_PATTERNS)


def escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def merge_orphan_bullets(lines: list[dict]) -> list[dict]:
    """Linje med kun ett tegn (bullet-symbol eller losen bokstav) merges med naeste tekstlinje.

    PDF-er har ofte bullets på egen linje fordi de er hevet/lavet litt vs. teksten, slik at
    extract_words havner i ulik y-bucket. Resultat: <p>•</p> over <p>tekst</p>. Vi slaar sammen.
    """
    out: list[dict] = []
    i = 0
    while i < len(lines):
        text = lines[i]["text"].strip()
        is_bullet = text in BULLET_CHARS or len(text) == 1
        next_is_long = i + 1 < len(lines) and len(lines[i + 1]["text"].strip()) > 2
        if is_bullet and next_is_long:
            merged = dict(lines[i + 1])
            merged["text"] = f"{text} {merged['text']}"
            out.append(merged)
            i += 2
        else:
            out.append(lines[i])
            i += 1
    return out


def filter_vertical_lines(lines: list[dict]) -> list[dict]:
    """Drop kluster av 3+ paafoelgende linjer med 1-2 tegn (vertikal sidebar-tekst)."""
    keep = [True] * len(lines)
    i = 0
    while i < len(lines):
        j = i
        while j < len(lines) and len(lines[j]["text"].strip()) <= 2:
            j += 1
        if j - i >= 3:
            for k in range(i, j):
                keep[k] = False
            i = j
        else:
            i = max(i + 1, j)
    return [ln for ln, k in zip(lines, keep) if k]


def group_lines_into_paragraphs(lines: list[dict]) -> list[dict]:
    """Slaa sammen paafoelgende linjer til paragrafer naar layouten tilsier det.

    Bryt paragraf ved: stort y-gap, font-stoerrelse-endring, eller setningsslutt fulgt av
    ny linje som starter med stor bokstav.
    """
    if not lines:
        return []
    paragraphs: list[dict] = []
    current = {"text": lines[0]["text"].strip(), "size": lines[0]["size"], "top": lines[0]["top"]}
    prev_bottom = lines[0]["bottom"]
    prev_size = lines[0]["size"]

    for ln in lines[1:]:
        text = ln["text"].strip()
        if not text:
            continue
        gap = ln["top"] - prev_bottom
        size_diff = abs(ln["size"] - prev_size) > 0.5
        line_height = max(prev_size, ln["size"])
        big_gap = gap > line_height * PARAGRAPH_GAP_MULTIPLIER

        if size_diff or big_gap:
            paragraphs.append(current)
            current = {"text": text, "size": ln["size"], "top": ln["top"]}
        else:
            current["text"] = _join_with_hyphenation(current["text"], text)
        prev_bottom = ln["bottom"]
        prev_size = ln["size"]
    paragraphs.append(current)
    return paragraphs


def _join_with_hyphenation(a: str, b: str) -> str:
    if HYPHEN_END.search(a):
        stripped = re.sub(r"[\-­‐‑]\s*$", "", a)
        if b and b[0].islower():
            return stripped + b
        return stripped + b
    return a + " " + b


def classify(paragraph: dict, body_size: float) -> str:
    sz = paragraph["size"]
    if sz >= body_size * 1.4:
        return "h2"
    if sz >= body_size * 1.2:
        return "h3"
    if sz >= body_size * 1.08:
        return "h4"
    return "p"


def sentence_case(text: str) -> str:
    """Lowercase alt, saa kapitaliser foerste alfabetiske tegn."""
    lower = text.lower()
    for i, ch in enumerate(lower):
        if ch.isalpha():
            return lower[:i] + ch.upper() + lower[i + 1 :]
    return text


_HEADING_NUMBER_LOWER = re.compile(r"^\d+(\.\d+)*\s+[a-zåäöæøü]", re.UNICODE)


def needs_heading_normalization(text: str) -> bool:
    """Sjekk om heading-tekst er feilformatert pga smaaskaape eller dårlig case.

    Triggrer paa:
    1. Mid-word caps i mixed-case ord (>= 1) — typisk "asiaKaslÄHtÖisYYs"
    2. Heading starter med tall + space + lowercase bokstav — typisk "2 suunnittelu"

    Brukes KUN paa headings (h2/h3/h4), ikke p — forkortelser som "UPM-Kymmene" eller
    "URN:NBN:..." i broedtekst skal ikke nedsettes.
    """
    weird = 0
    for word in text.split():
        if word.isupper() or word.islower():
            continue
        for i, ch in enumerate(word):
            if i > 0 and ch.isupper():
                weird += 1
    if weird >= 1:
        return True
    return bool(_HEADING_NUMBER_LOWER.match(text))


def normalize_heading(text: str) -> str:
    if needs_heading_normalization(text):
        return sentence_case(text)
    for i, ch in enumerate(text):
        if ch.isalpha():
            if ch.islower():
                return text[:i] + ch.upper() + text[i + 1 :]
            break
    return text


def render_table(rows: list[list[str | None]]) -> str:
    parts = ["<table><tbody>"]
    for row in rows:
        parts.append("<tr>")
        for cell in row:
            parts.append(f"<td>{escape((cell or '').strip())}</td>")
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "".join(parts)


_TABLE_CAPTION = re.compile(r"^\s*(Taulukko|Tabell|Table|Tabel)\s+\d+\b", re.I)


def _detect_caption_table_regions(page) -> list[tuple[float, float, str]]:
    """Finn (top, bottom, name) for tabell-regioner identifisert via 'Taulukko N.'-caption.

    Strategi: finn linjer som starter med caption-paatern, og capture fra caption til
    neste 'large gap' (> 30 punkt) eller heading-stoerrelse-skifte.
    """
    out: list[tuple[float, float, str]] = []
    try:
        words = page.extract_words(extra_attrs=["size"])
    except Exception:
        return out
    if not words:
        return out
    by_y: dict[int, list[dict]] = {}
    for w in words:
        y = round(float(w["top"]))
        by_y.setdefault(y, []).append(w)
    sorted_ys = sorted(by_y.keys())

    for i, y in enumerate(sorted_ys):
        line_words = by_y[y]
        line_text = " ".join(w["text"] for w in line_words).strip()
        if not _TABLE_CAPTION.match(line_text):
            continue
        top = float(y)
        bottom = float(page.height)
        for j in range(i + 1, len(sorted_ys)):
            ny = sorted_ys[j]
            gap = ny - sorted_ys[j - 1]
            if gap > 30:
                bottom = float(sorted_ys[j - 1]) + 15
                break
        m = re.match(r"^\s*(Taulukko|Tabell|Table|Tabel)\s+(\d+)", line_text, re.I)
        num = m.group(2) if m else "X"
        out.append((top, bottom, f"caption_table_{num}"))
    return out


def _bbox_overlap_ratio(a: tuple, b: tuple) -> float:
    """Returner overlap-ratio mellom to bbox-er (Intersection over min-area)."""
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0, iy0 = max(ax0, bx0), max(ay0, by0)
    ix1, iy1 = min(ax1, bx1), min(ay1, by1)
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    inter = (ix1 - ix0) * (iy1 - iy0)
    min_area = min((ax1 - ax0) * (ay1 - ay0), (bx1 - bx0) * (by1 - by0))
    return inter / min_area if min_area > 0 else 0.0


def extract_page_tables(
    page,
    pdf_path: Path | None = None,
    page_num: int = 0,
    image_dir: Path | None = None,
) -> list[tuple[float, float, str]]:
    """Returner liste av (top_y, bottom_y, html) for hver tabell paa siden.

    Bruker BAADE lines-baseret OG text-baseret detection (sistnevnte fanger tabeller uten
    synlige linjer — som Korjuun Taulukko 1). Deduplisere overlappende detections.
    Filtrer ut dekorative single-cell-bokser. Komplekse tabeller rendres som bilde.
    """
    out: list[tuple[float, float, str]] = []
    try:
        all_tables = list(page.find_tables())
    except Exception:
        all_tables = []

    caption_regions = _detect_caption_table_regions(page) if (pdf_path and image_dir) else []
    for top, bottom, name in caption_regions:
        filename = render_table_as_image(
            pdf_path, page_num,
            (0, top, float(page.width), bottom),
            image_dir, f"page{page_num}_{name}",
        )
        if filename:
            out.append((top, bottom, f'<p>[ TABLE — attachment: {filename} — page {page_num} ]</p>'))

    for idx, table in enumerate(all_tables, start=1):
        bbox = table.bbox
        if any(top <= bbox[1] <= bottom or top <= bbox[3] <= bottom for top, bottom, _ in caption_regions):
            continue
        try:
            rows = table.extract()
        except Exception:
            continue

        if not rows:
            continue
        non_empty_rows = [r for r in rows if any((c or "").strip() for c in r)]
        if len(non_empty_rows) < 2:
            continue
        max_cols = max((sum(1 for c in r if (c or "").strip()) for r in non_empty_rows), default=0)
        if max_cols < 2:
            continue
        bbox = table.bbox
        complex = is_complex_table(rows)
        rendered = None
        if complex and pdf_path and image_dir:
            filename = render_table_as_image(pdf_path, page_num, bbox, image_dir, f"page{page_num}_table{idx}")
            if filename:
                rendered = f'<p>[ TABLE — attachment: {filename} — page {page_num} ]</p>'
        if rendered is None:
            rendered = render_table(rows)
        out.append((float(bbox[1]), float(bbox[3]), rendered))
    return out


def is_toc(text: str) -> bool:
    """Detekter TOC-paragraf. Krev faktiske bokstav-ord (ikke kun tall + tankestrek),
    ellers fanger vi tabell-data som '70 - 24 - 80 - ...' falskt.
    """
    words = text.split()
    if len(words) < 2:
        return False
    alpha_words = sum(1 for w in words if any(c.isalpha() for c in w))
    if len(words) >= 10:
        numeric = sum(1 for w in words if w.replace(".", "").isdigit() and any(c.isdigit() for c in w))
        return (
            numeric >= 5
            and numeric / len(words) >= 0.2
            and alpha_words / len(words) >= 0.3
        )
    last = words[-1]
    if last.isdigit() and 1 <= int(last) <= 999:
        non_digit = sum(1 for w in words[:-1] if not w.replace(".", "").isdigit())
        return non_digit >= 1 and alpha_words >= 1
    return False


def filter_diagram_clusters(elements: list[tuple[str, str]]) -> tuple[list[tuple[str, str]], int]:
    """Drop kluster av 3+ paafoelgende korte (<=25 tegn) p-elementer.

    PDF-diagrammer (Kuva 1 osv.) har korte tekst-labels rundt seg som blir egne paragrafer.
    Et kluster av slike er sannsynligvis diagram-elementer.
    """
    out: list[tuple[str, str]] = []
    dropped = 0
    i = 0
    while i < len(elements):
        j = i
        while j < len(elements) and elements[j][0] == "p" and len(elements[j][1].strip()) <= 25:
            j += 1
        if j - i >= 3:
            dropped += j - i
            i = j
        else:
            out.append(elements[i])
            i += 1
    return out, dropped


def is_complex_table(rows: list[list[str | None]]) -> bool:
    """True hvis tabellen er kompleks (mange kolonner/rader/tomme celler) — heller render som bilde."""
    if not rows:
        return False
    cols = max((len(r) for r in rows), default=0)
    if cols > 4 or len(rows) > 15:
        return True
    total = sum(len(r) for r in rows)
    empty = sum(1 for r in rows for c in r if not (c or "").strip())
    return total > 0 and empty / total > 0.4


def render_table_as_image(pdf_path: Path, page_num: int, bbox: tuple, image_dir: Path, name: str) -> str | None:
    """Render bbox-region av PDF-side som PNG til image_dir. Returner filnavn."""
    try:
        import fitz
    except ImportError:
        return None
    doc = fitz.open(str(pdf_path))
    try:
        page = doc[page_num - 1]
        rect = fitz.Rect(*bbox)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=rect)
        filename = f"{name}.png"
        image_dir.mkdir(parents=True, exist_ok=True)
        pix.save(str(image_dir / filename))
        return filename
    except Exception:
        return None
    finally:
        doc.close()


_CHAPTER_PREFIX = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+", re.UNICODE)
_SUB_CHAPTER_PREFIX = re.compile(r"^(\d+\.\d+(?:\.\d+)*)\.?\s+", re.UNICODE)
_TOC_ITEM = re.compile(r"(\d+(?:\.\d+)*)\s+([^\d]+?(?:\s+[^\d]+)*?)\s+\d+(?=\s|\Z)", re.UNICODE)


def parse_toc_items(text: str) -> list[tuple[str, str]]:
    """Returner liste av (chapter_number, title). Krev at title har minst en bokstav,
    ellers er det sannsynligvis tabell-data med tankestrek (ikke ekte TOC).
    """
    return [
        (m.group(1), m.group(2).strip())
        for m in _TOC_ITEM.finditer(text)
        if any(c.isalpha() for c in m.group(2))
    ]


def heading_chapter(text: str) -> str | None:
    """Hvis text starter med 'N' eller 'N.N' etc., returner det. Ellers None."""
    m = _CHAPTER_PREFIX.match(text)
    return m.group(1) if m else None


_INLINE_SUB_CHAPTER = re.compile(
    r"\b(\d+\.\d+(?:\.\d+)*)\s+([^\s.]+(?:\s+[^\s.]+){0,3})\s+(?=[A-ZÅÄÖÆØÜ])",
    re.UNICODE,
)


def promote_chapter_paragraphs(elements: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Promoter korte 'N.N tekst'-paragrafer til h3/h4 + splitt p-er med inline kapittel-numre.

    To strategier:
    1. Kort p (<=80 tegn) som starter med 'N.N' eller 'N.N.N' -> h3 eller h4 direkte.
    2. Lang p som inneholder inline 'N.N tittel CapitalNext'-moenster -> splitt:
       (eventuell foran-tekst som p) + (heading) + (rest som p).

    Strategi 2 fanger PDF-er der underkapittel-overskriften ble joinet inn i broedtekst.
    """
    out: list[tuple[str, str]] = []
    for tag, text in elements:
        if tag != "p":
            out.append((tag, text))
            continue

        if len(text.strip()) <= 80:
            m = _SUB_CHAPTER_PREFIX.match(text)
            if m:
                chapter = m.group(1)
                dots = chapter.count(".")
                out.append(("h3" if dots == 1 else "h4", text))
                continue

        matches = list(_INLINE_SUB_CHAPTER.finditer(text))
        if not matches:
            out.append((tag, text))
            continue

        last_end = 0
        for m in matches:
            chapter, title = m.group(1), m.group(2)
            before = text[last_end:m.start()].strip()
            if before:
                out.append(("p", before))
            dots = chapter.count(".")
            out.append(("h3" if dots == 1 else "h4", f"{chapter} {title}"))
            last_end = m.end()
        rest = text[last_end:].strip()
        if rest:
            out.append(("p", rest))
    return out


def build_nested_toc(items: list[tuple[str, str]], chapter_to_id: dict[str, str]) -> str:
    """Bygg nested <ul> basert paa dot-count i hvert kapittel-nummer."""
    if not items:
        return ""

    nodes: list[tuple[int, str]] = []
    for chapter, title in items:
        depth = chapter.count(".")
        target = chapter_to_id.get(chapter)
        label = f"{chapter} {title}"
        link = f'<a href="#{target}">{escape(label)}</a>' if target else escape(label)
        nodes.append((depth, link))

    def render(start_idx: int, this_depth: int) -> tuple[str, int]:
        html = ["<ul>"]
        idx = start_idx
        while idx < len(nodes):
            depth, link = nodes[idx]
            if depth < this_depth:
                break
            html.append(f"<li>{link}")
            if idx + 1 < len(nodes) and nodes[idx + 1][0] > depth:
                sub, idx = render(idx + 1, depth + 1)
                html.append(sub)
            else:
                idx += 1
            html.append("</li>")
        html.append("</ul>")
        return "".join(html), idx

    result, _ = render(0, 0)
    return result


def post_process(elements: list[tuple[str, str]]) -> tuple[list[tuple[str, str]], int]:
    """Normaliser heading-case, generer heading-id, konverter TOC til nested <ul> med ankere."""
    elements = promote_chapter_paragraphs(elements)
    elements, _diagram_dropped = filter_diagram_clusters(elements)

    chapter_to_id: dict[str, str] = {}
    intermediate: list[tuple[str, str]] = []
    counter = 0

    for tag, text in elements:
        if tag == "__raw__":
            intermediate.append((tag, text))
            continue
        if tag in ("h2", "h3", "h4"):
            normalized = normalize_heading(text)
            counter += 1
            slug = f"sec-{counter}"
            chapter = heading_chapter(normalized)
            if chapter:
                chapter_to_id[chapter] = slug
            heading_html = f'<{tag} id="{slug}">{escape(normalized)}</{tag}>'
            intermediate.append(("__raw__", heading_html))
        else:
            intermediate.append((tag, text))

    final: list[tuple[str, str]] = []
    toc_dropped = 0
    for tag, text in intermediate:
        if tag == "__raw__":
            final.append((tag, text))
            continue
        if tag == "p" and is_toc(text):
            items = parse_toc_items(text)
            if items:
                final.append(("__raw__", build_nested_toc(items, chapter_to_id)))
                continue
        final.append((tag, text))
    return final, toc_dropped


def split_bullets(text: str) -> tuple[str, list[str]]:
    """Returner (lead_text, [item1, item2, ...]). Tom items-liste hvis ingen bullets."""
    parts = BULLET_SPLIT.split(text)
    if len(parts) < 2:
        return text, []
    lead = parts[0].strip()
    items = [p.strip() for p in parts[1:] if p.strip()]
    return lead, items


def wrap_bullet_lists(elements: list[tuple[str, str]]) -> list[str]:
    """Konverter <p>... • a • b • c</p> til <p>lead</p><ul><li>a</li>...</ul>.

    "__raw__"-tag er ferdig HTML (f.eks. tabeller) og slippes gjennom uten escape.
    Slaar ogsaa sammen losene <li>-elementer (fra DOCX) til <ul>.
    """
    out: list[str] = []
    pending: list[str] = []

    def flush() -> None:
        if pending:
            items_html = "".join(f"<li>{escape(item)}</li>" for item in pending)
            out.append(f"<ul>{items_html}</ul>")
            pending.clear()

    for tag, text in elements:
        if tag == "__raw__":
            flush()
            out.append(text)
            continue
        if tag == "li":
            pending.append(text)
            continue
        if tag == "p":
            lead, items = split_bullets(text)
            if items:
                if lead:
                    flush()
                    out.append(f"<p>{escape(lead)}</p>")
                pending.extend(items)
                continue
        flush()
        out.append(f"<{tag}>{escape(text)}</{tag}>")
    flush()
    return out


def validate_and_dedupe(elements: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Drop for korte headings, dedup gjentatte headings, slaa sammen paafoelgende same-level headings."""
    out: list[tuple[str, str]] = []
    recent_headings: list[str] = []

    for tag, text in elements:
        if tag == "__raw__":
            out.append((tag, text))
            continue
        if tag in ("h2", "h3", "h4"):
            stripped = text.strip()
            if len(stripped) < MIN_HEADING_LEN:
                continue
            if len(stripped) > MAX_HEADING_LEN:
                out.append(("p", text))
                continue
            key = stripped.lower()
            if key in recent_headings:
                continue
            if out and out[-1][0] == tag and len(out[-1][1]) <= HEADING_MERGE_MAX_LEN and len(text) <= HEADING_MERGE_MAX_LEN:
                merged_text = out[-1][1].rstrip() + " " + text.lstrip()
                out[-1] = (tag, merged_text)
                recent_headings.append(merged_text.strip().lower())
            else:
                out.append((tag, text))
                recent_headings.append(key)
            if len(recent_headings) > HEADING_DEDUPE_WINDOW:
                recent_headings = recent_headings[-HEADING_DEDUPE_WINDOW:]
        else:
            out.append((tag, text))
    return out


def convert_pdf(path: Path, image_dir: Path | None = None, embed_images: bool = False) -> tuple[str, list[str]]:
    import pdfplumber

    warnings: list[str] = []
    elements: list[tuple[str, str]] = []
    total_chars = 0
    image_count = 0
    table_count = 0

    extracted_images: list[dict] = []
    if image_dir is not None:
        try:
            extracted_images = extract_pdf_images(path, image_dir)
        except Exception as e:
            warnings.append(f"bilde-ekstraksjon feilet: {e}")
    images_by_page: dict[int, list[dict]] = {}
    for info in extracted_images:
        images_by_page.setdefault(info["page"], []).append(info)
    image_use_idx: dict[int, int] = {}

    with pdfplumber.open(path) as pdf:
        sizes: list[float] = []
        for page in pdf.pages:
            for ch in page.chars:
                sizes.append(round(ch.get("size", 10.0), 1))
        body_size = Counter(sizes).most_common(1)[0][0] if sizes else 10.0
        page_count = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            page_elements: list[tuple[float, str, str]] = []

            page_tables = extract_page_tables(page, pdf_path=path, page_num=page_num, image_dir=image_dir)
            for top_y, _bot_y, table_html in page_tables:
                page_elements.append((top_y, "__raw__", table_html))
                table_count += 1

            for img in page.images:
                image_count += 1
                top_y = float(img.get("top", 0) or 0)
                filename = None
                pool = images_by_page.get(page_num, [])
                if pool:
                    used = image_use_idx.get(page_num, 0)
                    if used < len(pool):
                        filename = pool[used]["filename"]
                        image_use_idx[page_num] = used + 1
                tag, content = image_element(page_num, filename, image_dir, embed=embed_images)
                page_elements.append((top_y, tag, content))

            try:
                words = page.extract_words(extra_attrs=["size"])
            except Exception as e:
                warnings.append(f"side {page_num}: extract_words feilet ({e})")
                words = []

            image_bboxes = [
                (float(img.get("top", 0) or 0), float(img.get("bottom", 0) or 0),
                 float(img.get("x0", 0) or 0), float(img.get("x1", 0) or 0))
                for img in page.images
            ]
            if page_tables or image_bboxes:
                def _in_table(w):
                    return any(top <= float(w["top"]) <= bot for top, bot, _ in page_tables)
                def _in_image(w):
                    wtop, wx0 = float(w["top"]), float(w["x0"])
                    return any(itop <= wtop <= ibot and ix0 <= wx0 <= ix1 for itop, ibot, ix0, ix1 in image_bboxes)
                words = [w for w in words if not _in_table(w) and not _in_image(w)]

            if not words:
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    if is_noise(line) or not line.strip():
                        continue
                    page_elements.append((0.0, "p", line.strip()))
                    total_chars += len(line)
            else:
                by_y: dict[int, list[dict]] = {}
                for w in words:
                    key = round(float(w["top"]))
                    by_y.setdefault(key, []).append(w)

                page_lines: list[dict] = []
                for y in sorted(by_y.keys()):
                    line_words = by_y[y]
                    text = " ".join(w["text"] for w in line_words).strip()
                    if not text or is_noise(text):
                        continue
                    sz = max((float(w.get("size", 10.0)) for w in line_words), default=10.0)
                    bottom = max(float(w.get("bottom", y)) for w in line_words)
                    page_lines.append({"text": text, "size": sz, "top": float(y), "bottom": bottom})

                page_lines = filter_vertical_lines(page_lines)
                page_lines = merge_orphan_bullets(page_lines)
                paragraphs = group_lines_into_paragraphs(page_lines)

                for para in paragraphs:
                    tag = classify(para, body_size)
                    page_elements.append((para["top"], tag, para["text"]))
                    total_chars += len(para["text"])

            page_elements.sort(key=lambda e: e[0])
            for _, tag, content in page_elements:
                elements.append((tag, content))

    if image_count:
        warnings.append(f"{image_count} bilde(r) ble erstattet med placeholder")
    if table_count:
        warnings.append(f"{table_count} tabell(er) ekstrahert")
    if not table_count and page_count and total_chars / page_count < SCANNED_PDF_CHARS_PER_PAGE:
        warnings.append(
            f"PDF inneholder kun {total_chars} tegn over {page_count} sider "
            f"({total_chars // max(page_count,1)}/side) — sannsynligvis skannet. "
            f"Vurder OCR (Tesseract eller vision-modell)."
        )

    elements = validate_and_dedupe(elements)
    elements, toc_dropped = post_process(elements)
    if toc_dropped:
        warnings.append(f"{toc_dropped} TOC-paragraf(er) droppet")
    html = "\n".join(wrap_bullet_lists(elements))
    return html, warnings


def _docx_inner(runs, text_fallback: str) -> str:
    parts: list[str] = []
    for run in runs:
        t = escape(run.text)
        if not t:
            continue
        if run.bold:
            t = f"<strong>{t}</strong>"
        if run.italic:
            t = f"<em>{t}</em>"
        parts.append(t)
    return "".join(parts) if parts else escape(text_fallback)


def _detect_docx_heading_level(text: str, runs) -> str:
    """Returner h2/h3/h4 hvis paragrafen er all-bold + heading-aktig. Ellers 'p'."""
    if not runs or len(text) > 80:
        return "p"
    runs_with_text = [r for r in runs if r.text.strip()]
    if not runs_with_text:
        return "p"
    if not all(r.bold for r in runs_with_text):
        return "p"
    m = _CHAPTER_PREFIX.match(text)
    if m:
        dots = m.group(1).count(".")
        return "h3" if dots == 0 else "h4"
    alpha = [c for c in text if c.isalpha()]
    if alpha and all(c.isupper() for c in alpha):
        return "h2"
    return "p"


def convert_docx(path: Path, image_dir: Path | None = None, embed_images: bool = False) -> tuple[str, list[str]]:
    import docx

    doc = docx.Document(path)
    html_parts: list[str] = []
    warnings: list[str] = []

    extracted_images: list[dict] = []
    if image_dir is not None:
        try:
            extracted_images = extract_docx_images(path, image_dir)
        except Exception as e:
            warnings.append(f"bilde-ekstraksjon feilet: {e}")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()

        if "heading 1" in style:
            tag = "h2"
        elif "heading 2" in style:
            tag = "h3"
        elif "heading 3" in style or "heading 4" in style:
            tag = "h4"
        elif "list" in style:
            tag = "li"
        else:
            tag = _detect_docx_heading_level(text, para.runs)

        if tag in ("h2", "h3", "h4"):
            inner = escape(text)
        elif tag == "p" and _SUB_CHAPTER_PREFIX.match(text) and not any(r.bold for r in para.runs):
            m = _SUB_CHAPTER_PREFIX.match(text)
            prefix_text = text[: m.end()].rstrip()
            rest = text[m.end():]
            inner = f"<strong>{escape(prefix_text)}</strong> {escape(rest)}"
        else:
            inner = _docx_inner(para.runs, text)
        html_parts.append(f"<{tag}>{inner}</{tag}>")

    for table in doc.tables:
        html_parts.append("<table>")
        html_parts.append("<tbody>")
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                html_parts.append(f"<td>{escape(cell.text.strip())}</td>")
            html_parts.append("</tr>")
        html_parts.append("</tbody>")
        html_parts.append("</table>")

    image_count = len(doc.inline_shapes)
    for i in range(image_count):
        filename = extracted_images[i]["filename"] if i < len(extracted_images) else None
        tag, content = image_element(None, filename, image_dir, embed=embed_images)
        if tag == "__raw__":
            html_parts.append(content)
        else:
            html_parts.append(f"<p>{content}</p>")
    if image_count:
        warnings.append(f"{image_count} bilde(r) prosessert")
    if extracted_images:
        warnings.append(f"{len(extracted_images)} bilde(r) lagret i {image_dir}")

    return "\n".join(html_parts), warnings


def sanitize(html: str) -> str:
    try:
        import bleach
    except ImportError:
        print("WARN: bleach ikke installert — hopper over rensing", file=sys.stderr)
        return html
    return bleach.clean(
        html,
        tags=ALLOWED_TAGS,
        attributes=ALLOWED_ATTRS,
        protocols=ALLOWED_PROTOCOLS,
        strip=True,
    )


def main() -> int:
    ap = argparse.ArgumentParser(description="Konverter PDF/DOCX til CKEditor-vennlig HTML")
    ap.add_argument("input", type=Path, help="PDF eller DOCX")
    ap.add_argument("-o", "--output", type=Path, help="Skriv output til fil (default stdout)")
    ap.add_argument("--no-images", action="store_true", help="Hopp over bilde-ekstraksjon")
    ap.add_argument("--embed-images", action="store_true", help="Inline bilder som base64 i HTML (stor output)")
    args = ap.parse_args()

    if not args.input.exists():
        print(f"FEIL: finner ikke {args.input}", file=sys.stderr)
        return 1

    image_dir: Path | None = None
    if args.output and not args.no_images:
        image_dir = args.output.parent / f"{args.output.stem}_bilder"

    ext = args.input.suffix.lower()
    if ext == ".pdf":
        html, warnings = convert_pdf(args.input, image_dir=image_dir, embed_images=args.embed_images)
    elif ext == ".docx":
        html, warnings = convert_docx(args.input, image_dir=image_dir, embed_images=args.embed_images)
    else:
        print(f"FEIL: ustøttet format {ext}", file=sys.stderr)
        return 1

    html = sanitize(html)

    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(html, encoding="utf-8")
        print(f"Skrev {len(html)} tegn til {args.output}", file=sys.stderr)
    else:
        sys.stdout.write(html)
        sys.stdout.write("\n")

    for w in warnings:
        print(f"[warn] {w}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
